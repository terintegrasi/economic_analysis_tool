import streamlit as st
import pandas as pd
import numpy as np
from utils.data_processor import DataProcessor
from utils.economic_analyzer import EconomicAnalyzer
from utils.visualizer import Visualizer

# Page configuration
st.set_page_config(
    page_title="SINERGI-P",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

def main():
    st.title("📊 SINERGI-P")
    st.markdown("""
    This tool analyzes expenditure GDP data to identify which types of expenditure are driving economic growth in your region.
    Upload your CSV file containing expenditure GDP data to get started.
    """)
    
    # Initialize session state
    if 'data_uploaded' not in st.session_state:
        st.session_state.data_uploaded = False
    if 'analysis_complete' not in st.session_state:
        st.session_state.analysis_complete = False
    
    # Sidebar for file upload and settings
    with st.sidebar:
        st.header("📁 Data Upload")
        uploaded_file = st.file_uploader(
            "Upload your CSV file with expenditure GDP data",
            type=['csv'],
            help="Upload a CSV file containing time series data of GDP expenditure components"
        )
        
        if uploaded_file is not None:
            try:
                # Process the uploaded file
                processor = DataProcessor()
                df = processor.load_and_validate_csv(uploaded_file)
                
                if df is not None:
                    st.success("✅ File uploaded successfully!")
                    st.session_state.data_uploaded = True
                    st.session_state.df = df
                    
                    # Display basic info about the dataset
                    st.subheader("📋 Dataset Info")
                    st.write(f"**Rows:** {len(df)}")
                    st.write(f"**Columns:** {len(df.columns)}")
                    st.write(f"**Date Range:** {df.index.min()} to {df.index.max()}")
                    
                    # Column mapping interface
                    st.subheader("🔗 Column Mapping")
                    st.write("Map your CSV columns to expenditure categories:")
                    
                    available_columns = df.columns.tolist()
                    
                    consumption_col = st.selectbox(
                        "Consumption Expenditure",
                        options=['None'] + available_columns,
                        help="Select the column representing household/private consumption"
                    )
                    
                    investment_col = st.selectbox(
                        "Investment Expenditure", 
                        options=['None'] + available_columns,
                        help="Select the column representing gross fixed capital formation/investment"
                    )
                    
                    government_col = st.selectbox(
                        "Government Expenditure",
                        options=['None'] + available_columns,
                        help="Select the column representing government spending"
                    )
                    
                    exports_col = st.selectbox(
                        "Exports",
                        options=['None'] + available_columns,
                        help="Select the column representing exports"
                    )
                    
                    imports_col = st.selectbox(
                        "Imports",
                        options=['None'] + available_columns,
                        help="Select the column representing imports"
                    )
                    
                    gdp_col = st.selectbox(
                        "Total GDP",
                        options=['None'] + available_columns,
                        help="Select the column representing total GDP (optional - can be calculated)"
                    )
                    
                    # Analysis settings
                    st.subheader("⚙️ Analysis Settings")
                    
                    # Primary Growth Driver Analysis Method
                    driver_analysis_method = st.selectbox(
                        "Primary Growth Driver Analysis Method",
                        options=['correlation_contribution', 'log_log_regression', 'average_elasticity'],
                        format_func=lambda x: {
                            'correlation_contribution': 'Correlation & Contribution (Default)',
                            'log_log_regression': 'Log-Log Regression Analysis',
                            'average_elasticity': 'Average Elasticity Method'
                        }[x],
                        help="Choose the method to identify the primary economic growth driver"
                    )
                    
                    analysis_period = st.selectbox(
                        "Analysis Period",
                        options=['Full Dataset', 'Last 5 Years', 'Last 10 Years', 'Custom'],
                        help="Select the time period for analysis"
                    )
                    
                    if analysis_period == 'Custom':
                        start_date = st.date_input("Start Date", value=df.index.min())
                        end_date = st.date_input("End Date", value=df.index.max())
                    
                    # Forecasting settings
                    st.subheader("🔮 Forecasting Settings")
                    enable_forecasting = st.checkbox("Enable Economic Forecasting", value=False)
                    
                    if enable_forecasting:
                        forecast_method = st.selectbox(
                            "Forecast Method",
                            options=[
                                'linear_regression',
                                'polynomial', 
                                'moving_average',
                                'random_forest',
                                'svm',
                                'var_model',
                                'gmm',
                                'cge_model'
                            ],
                            format_func=lambda x: {
                                'linear_regression': 'Linear Regression',
                                'polynomial': 'Polynomial Regression',
                                'moving_average': 'Moving Average',
                                'random_forest': 'Random Forest',
                                'svm': 'Support Vector Machine',
                                'var_model': 'Vector Autoregression (VAR)',
                                'gmm': 'Generalized Method of Moments (GMM)',
                                'cge_model': 'Computable General Equilibrium (CGE)'
                            }[x],
                            help="Select the forecasting method to use"
                        )
                        
                        forecast_years = st.slider(
                            "Forecast Years",
                            min_value=1,
                            max_value=10,
                            value=5,
                            help="Number of years to forecast into the future"
                        )
                        
                        st.session_state.enable_forecasting = True
                        st.session_state.forecast_method = forecast_method
                        st.session_state.forecast_years = forecast_years
                    else:
                        st.session_state.enable_forecasting = False
                    
                    # Store column mappings in session state
                    column_mapping = {
                        'consumption': consumption_col if consumption_col != 'None' else None,
                        'investment': investment_col if investment_col != 'None' else None,
                        'government': government_col if government_col != 'None' else None,
                        'exports': exports_col if exports_col != 'None' else None,
                        'imports': imports_col if imports_col != 'None' else None,
                        'gdp': gdp_col if gdp_col != 'None' else None
                    }
                    
                    st.session_state.column_mapping = column_mapping
                    
                    # Run analysis button
                    if st.button("🔍 Run Economic Analysis", type="primary"):
                        if any(col is not None for col in column_mapping.values()):
                            st.session_state.analysis_complete = True
                            st.rerun()
                        else:
                            st.error("Please map at least one expenditure column to run the analysis.")
                
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                st.session_state.data_uploaded = False
    
    # Main content area
    if st.session_state.data_uploaded and st.session_state.analysis_complete:
        # Run the economic analysis
        df = st.session_state.df
        column_mapping = st.session_state.column_mapping
        
        # Initialize analyzer and visualizer
        analyzer = EconomicAnalyzer()
        visualizer = Visualizer()
        
        # Prepare data for analysis
        analysis_data = analyzer.prepare_analysis_data(df, column_mapping)
        
        if analysis_data is not None and not analysis_data.empty:
            # Perform economic analysis
            results = analyzer.analyze_growth_drivers(analysis_data, method=driver_analysis_method)
            
            # Perform forecasting if enabled
            forecast_results = None
            accuracy_results = None
            if st.session_state.get('enable_forecasting', False):
                forecast_method = st.session_state.get('forecast_method', 'linear_regression')
                forecast_years = st.session_state.get('forecast_years', 5)
                
                with st.spinner(f"Running {forecast_method} forecasting for {forecast_years} years..."):
                    forecast_results = analyzer.forecast_growth_drivers(
                        analysis_data, 
                        forecast_years=forecast_years, 
                        method=forecast_method
                    )
                    
                    # Calculate forecast accuracy using backtesting
                    accuracy_results = analyzer.calculate_forecast_accuracy(
                        analysis_data, 
                        method=forecast_method
                    )
            
            # Display results
            display_analysis_results(results, analysis_data, visualizer, forecast_results, accuracy_results)
        else:
            st.error("Unable to prepare data for analysis. Please check your column mappings.")
    
    elif st.session_state.data_uploaded and not st.session_state.analysis_complete:
        st.info("📝 Please configure the column mappings in the sidebar and click 'Run Economic Analysis' to proceed.")
        
        # Show data preview
        if 'df' in st.session_state:
            st.subheader("📊 Data Preview")
            st.dataframe(st.session_state.df.head(10))
    
    else:
        # Show methodology and instructions
        display_methodology()

def display_analysis_results(results, data, visualizer, forecast_results=None, accuracy_results=None):
    """Display the complete analysis results"""
    
    st.header("📈 Economic Growth Analysis Results")
    
    # Key insights at the top
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "Primary Growth Driver",
            results['primary_driver'],
            help="The expenditure category with the highest contribution to GDP growth"
        )
    
    with col2:
        st.metric(
            "GDP Growth Rate",
            f"{results['gdp_growth_rate']:.2f}%",
            help="Average annual GDP growth rate over the analysis period"
        )
    
    with col3:
        st.metric(
            "Strongest Correlation",
            f"{results['strongest_correlation']:.3f}",
            help="Highest correlation coefficient between expenditure and GDP growth"
        )
    
    with col4:
        st.metric(
            "Analysis Period",
            f"{len(data)} periods",
            help="Number of time periods included in the analysis"
        )
    
    # Tabs for different analysis views
    tabs = [
        "📊 Growth Contributions", 
        "📈 Time Series Analysis", 
        "🔗 Correlation Analysis",
        "📋 Detailed Results",
        "💡 Methodology"
    ]
    
    # Add forecasting tabs if forecast results are available
    if forecast_results is not None:
        tabs.extend([
            "🔮 Economic Forecasts",
            "📊 Forecast Accuracy"
        ])
    
    tab_objects = st.tabs(tabs)
    tab1, tab2, tab3, tab4, tab5 = tab_objects[:5]
    
    if forecast_results is not None and len(tab_objects) >= 7:
        tab6, tab7 = tab_objects[5], tab_objects[6]
    else:
        tab6, tab7 = None, None
    
    with tab1:
        st.subheader("Growth Contributions by Expenditure Type")
        
        # Growth contributions chart
        if 'growth_contributions' in results:
            fig_contrib = visualizer.create_contribution_chart(results['growth_contributions'])
            st.plotly_chart(fig_contrib, use_container_width=True)
        
        # Summary table
        if 'summary_stats' in results:
            st.subheader("Summary Statistics")
            st.dataframe(results['summary_stats'])
    
    with tab2:
        st.subheader("Time Series Analysis")
        
        # Time series plots
        fig_ts = visualizer.create_time_series_plot(data)
        st.plotly_chart(fig_ts, use_container_width=True)
        
        # Growth rates over time
        if 'growth_rates' in results:
            fig_growth = visualizer.create_growth_rates_plot(results['growth_rates'])
            st.plotly_chart(fig_growth, use_container_width=True)
    
    with tab3:
        st.subheader("Correlation Analysis")
        
        # Correlation matrix
        if 'correlations' in results:
            fig_corr = visualizer.create_correlation_heatmap(results['correlations'])
            st.plotly_chart(fig_corr, use_container_width=True)
        
        # Method-specific analysis results
        analysis_method = results.get('analysis_method', 'correlation_contribution')
        
        if analysis_method == 'log_log_regression' and 'log_log_elasticities' in results:
            st.subheader("Log-Log Regression Elasticity Analysis")
            st.markdown("**📈 Elasticity coefficients from log-log regression (% change in GDP per 1% change in component)**")
            
            # Create log-log elasticity results table
            elasticity_table = []
            for category, elastic_result in results['log_log_elasticities'].items():
                category_name = {
                    'consumption': 'Consumption',
                    'investment': 'Investment',
                    'government': 'Government',
                    'exports': 'Exports',
                    'imports': 'Imports',
                    'net_exports': 'Net Exports'
                }.get(category, category.title())
                
                p_value = elastic_result.get('p_value', 1)
                is_significant = p_value < 0.05
                
                elasticity_table.append({
                    'Component': category_name,
                    'Elasticity': f"{elastic_result.get('elasticity', 0):.4f}",
                    'R-squared': f"{elastic_result.get('r_squared', 0):.4f}",
                    'P-value': f"{p_value:.4f}",
                    'Significance': 'Significant' if is_significant else 'Not Significant'
                })
            
            if elasticity_table:
                import pandas as pd
                elasticity_df = pd.DataFrame(elasticity_table)
                
                st.markdown("**Statistical significance threshold: p < 0.05**")
                
                # Create styled dataframe
                def color_p_values(val):
                    if 'P-value' in str(val):
                        return ''
                    try:
                        p_val = float(val)
                        if p_val < 0.05:
                            return 'color: green; font-weight: bold'
                        else:
                            return 'color: red; font-weight: bold'
                    except:
                        return ''
                
                # Apply styling to P-value column
                styled_df = elasticity_df.style.applymap(color_p_values, subset=['P-value'])
                st.dataframe(styled_df, use_container_width=True)
                
                st.info("📊 **Interpretation**: Elasticity > 1 indicates elastic response (proportionally larger GDP change), < 1 indicates inelastic response.")
        
        elif analysis_method == 'average_elasticity' and 'average_elasticities' in results:
            st.subheader("Average Elasticity Analysis")
            st.markdown("**📊 Average elasticity of GDP with respect to each expenditure component**")
            
            # Create average elasticity results table
            avg_elasticity_table = []
            for category, elastic_result in results['average_elasticities'].items():
                category_name = {
                    'consumption': 'Consumption',
                    'investment': 'Investment',
                    'government': 'Government',
                    'exports': 'Exports',
                    'imports': 'Imports',
                    'net_exports': 'Net Exports'
                }.get(category, category.title())
                
                elasticity_range = elastic_result.get('elasticity_range', {})
                
                avg_elasticity_table.append({
                    'Component': category_name,
                    'Average Elasticity': f"{elastic_result.get('average_elasticity', 0):.4f}",
                    'Absolute Avg Elasticity': f"{elastic_result.get('absolute_average_elasticity', 0):.4f}",
                    'Standard Deviation': f"{elastic_result.get('standard_deviation', 0):.4f}",
                    'Correlation': f"{elastic_result.get('correlation', 0):.4f}",
                    'Min/Max Range': f"{elasticity_range.get('min', 0):.3f} / {elasticity_range.get('max', 0):.3f}",
                    'Observations': str(elastic_result.get('observations', 0))
                })
            
            if avg_elasticity_table:
                import pandas as pd
                avg_elasticity_df = pd.DataFrame(avg_elasticity_table)
                st.dataframe(avg_elasticity_df, use_container_width=True)
                
                st.info("📊 **Interpretation**: Higher absolute average elasticity indicates stronger responsiveness of GDP to changes in that component. Positive correlation supports the relationship direction.")
        
        else:  # Default: correlation_contribution method
            # Regression analysis table
            if 'regression_results' in results:
                st.subheader("Regression Analysis Results")
                
                # Create regression results table
                regression_table = []
                for category, reg_result in results['regression_results'].items():
                    category_name = {
                        'consumption': 'Consumption',
                        'investment': 'Investment',
                        'government': 'Government',
                        'exports': 'Exports',
                        'imports': 'Imports',
                        'net_exports': 'Net Exports'
                    }.get(category, category.title())
                    
                    p_value = reg_result.get('p_value', 1)
                    is_significant = p_value < 0.05
                    
                    regression_table.append({
                        'Component': category_name,
                        'R-squared': f"{reg_result.get('r_squared', 0):.4f}",
                        'Coefficient': f"{reg_result.get('coefficient', 0):.4f}",
                        'P-value': f"{p_value:.4f}",
                        'Significance': 'Significant' if is_significant else 'Not Significant'
                    })
                
                if regression_table:
                    import pandas as pd
                    regression_df = pd.DataFrame(regression_table)
                    
                    # Display table with colored p-values
                    st.markdown("**Statistical significance threshold: p < 0.05**")
                    
                    # Create styled dataframe
                    def color_p_values(val):
                        if 'P-value' in str(val):
                            return ''
                        try:
                            p_val = float(val)
                            if p_val < 0.05:
                                return 'color: green; font-weight: bold'
                            else:
                                return 'color: red; font-weight: bold'
                        except:
                            return ''
                    
                    # Apply styling to P-value column
                    styled_df = regression_df.style.applymap(color_p_values, subset=['P-value'])
                    st.dataframe(styled_df, use_container_width=True)
    
    with tab4:
        st.subheader("Detailed Analysis Results")
        
        # Complete results in expandable sections
        with st.expander("📊 Growth Rate Analysis"):
            if 'growth_analysis' in results:
                st.json(results['growth_analysis'])
        
        with st.expander("🔢 Statistical Summary"):
            if 'statistical_summary' in results:
                st.dataframe(results['statistical_summary'])
        
        with st.expander("📈 Trend Analysis"):
            if 'trend_analysis' in results:
                st.write(results['trend_analysis'])
    
    with tab5:
        display_detailed_methodology()
    
    # Forecasting tabs (if forecast results are available)
    if forecast_results is not None and tab6 is not None:
        with tab6:
            st.subheader("Economic Forecasts")
            st.markdown("**📈 Future projections based on historical trends and selected forecasting method**")
            
            # Forecast method info
            forecast_method = st.session_state.get('forecast_method', 'linear_regression')
            forecast_years = st.session_state.get('forecast_years', 5)
            
            st.info(f"🔮 **Method**: {forecast_method.replace('_', ' ').title()} | **Forecast Period**: {forecast_years} years")
            
            # Growth scenario options
            st.subheader("📊 Growth Scenario Settings")
            growth_scenario = st.radio(
                "Choose forecast scenario:",
                options=['historical_pattern', 'custom_growth'],
                format_func=lambda x: {
                    'historical_pattern': 'Follow Historical Pattern',
                    'custom_growth': 'Custom Growth Rates'
                }[x],
                horizontal=True
            )
            
            custom_growth_rates = {}
            if growth_scenario == 'custom_growth':
                st.markdown("**Set custom annual growth rates for each component:**")
                
                # GDP Growth Rate setting
                gdp_growth_target = st.number_input(
                    "🎯 Target GDP Growth Rate (%)",
                    min_value=-20.0,
                    max_value=50.0,
                    value=4.0,
                    step=0.1,
                    help="Set desired GDP growth rate - other components will adjust automatically to achieve this target"
                )
                
                auto_adjust = st.checkbox(
                    "Auto-adjust other components to achieve GDP target",
                    value=True,
                    help="When enabled, other component growth rates will be calculated to meet the GDP target"
                )
                
                col1, col2 = st.columns(2)
                available_components = [comp for comp in forecast_results.keys() if comp in data.columns and comp != 'gdp']
                
                if auto_adjust:
                    st.info("💡 **Auto-adjustment enabled**: Component growth rates will be calculated to achieve the GDP target based on their historical contribution weights.")
                    
                    # Calculate historical contribution weights
                    if 'contributions' in results:
                        contributions = results['contributions']
                        total_contribution = sum([abs(contrib) for contrib in contributions.values() if contrib is not None])
                        
                        if total_contribution > 0:
                            for component in available_components:
                                if component in contributions and contributions[component] is not None:
                                    weight = abs(contributions[component]) / total_contribution
                                    # Adjust component growth rate proportionally
                                    adjusted_growth = gdp_growth_target * weight
                                    custom_growth_rates[component] = adjusted_growth
                                else:
                                    custom_growth_rates[component] = gdp_growth_target * 0.2  # Default weight
                        else:
                            # Equal distribution if no contribution data
                            equal_growth = gdp_growth_target / len(available_components)
                            for component in available_components:
                                custom_growth_rates[component] = equal_growth
                    else:
                        # Equal distribution fallback
                        equal_growth = gdp_growth_target / len(available_components)
                        for component in available_components:
                            custom_growth_rates[component] = equal_growth
                    
                    # Display calculated rates
                    st.markdown("**📊 Calculated Component Growth Rates:**")
                    for i, component in enumerate(available_components):
                        component_name = {
                            'consumption': 'Consumption',
                            'investment': 'Investment',
                            'government': 'Government',
                            'exports': 'Exports',
                            'imports': 'Imports'
                        }.get(component, component.title())
                        
                        with col1 if i % 2 == 0 else col2:
                            st.metric(
                                component_name,
                                f"{custom_growth_rates[component]:.2f}%",
                                help=f"Auto-calculated growth rate for {component_name}"
                            )
                
                else:
                    st.markdown("**Manual component growth rates:**")
                    for i, component in enumerate(available_components):
                        component_name = {
                            'consumption': 'Consumption',
                            'investment': 'Investment',
                            'government': 'Government',
                            'exports': 'Exports',
                            'imports': 'Imports'
                        }.get(component, component.title())
                        
                        with col1 if i % 2 == 0 else col2:
                            custom_growth_rates[component] = st.number_input(
                                f"{component_name} Growth Rate (%)",
                                min_value=-50.0,
                                max_value=100.0,
                                value=3.0,
                                step=0.1,
                                help=f"Annual growth rate for {component_name}"
                            )
                
                # Set GDP growth rate
                custom_growth_rates['gdp'] = gdp_growth_target
            
            # Individual component forecasts
            available_components = [comp for comp in forecast_results.keys() if comp in data.columns]
            
            if available_components:
                selected_component = st.selectbox(
                    "Select component to view detailed forecast:",
                    options=available_components,
                    format_func=lambda x: {
                        'consumption': 'Consumption',
                        'investment': 'Investment',
                        'government': 'Government',
                        'exports': 'Exports',
                        'imports': 'Imports',
                        'gdp': 'GDP'
                    }.get(x, x.title())
                )
                
                # Show individual forecast plot
                fig_forecast = visualizer.create_forecast_plot(data, forecast_results, selected_component)
                st.plotly_chart(fig_forecast, use_container_width=True)
                
                # Show forecast comparison
                st.subheader("Forecasted Growth Rate Comparison")
                fig_comparison = visualizer.create_forecast_comparison(forecast_results)
                st.plotly_chart(fig_comparison, use_container_width=True)
                
                # Forecast summary table with year-by-year breakdown
                st.subheader("Forecast Summary")
                forecast_summary = []
                
                for component, data_dict in forecast_results.items():
                    if 'predictions' in data_dict and len(data_dict['predictions']) > 0:
                        component_name = {
                            'consumption': 'Consumption',
                            'investment': 'Investment',
                            'government': 'Government',
                            'exports': 'Exports',
                            'imports': 'Imports',
                            'gdp': 'GDP'
                        }.get(component, component.title())
                        
                        # Create base row
                        row = {
                            'Component': component_name,
                            'Current Value': f"{data[component].iloc[-1]:,.0f}" if component in data.columns else None
                        }
                        
                        # Generate predictions based on scenario
                        if growth_scenario == 'custom_growth' and component in custom_growth_rates:
                            # Use custom growth rates
                            current_value = data[component].iloc[-1] if component in data.columns else 1000
                            predictions = []
                            for year in range(forecast_years):
                                future_value = current_value * ((1 + custom_growth_rates[component]/100) ** (year + 1))
                                predictions.append(future_value)
                        else:
                            # Use model predictions
                            predictions = data_dict['predictions']
                        
                        # Add forecast for each year
                        for year_idx in range(min(len(predictions), forecast_years)):
                            year_label = f'Year {year_idx + 1}'
                            row[year_label] = f"{predictions[year_idx]:,.0f}"
                        
                        # Add average annual growth
                        if component in data.columns and data[component].iloc[-1] != 0:
                            if growth_scenario == 'custom_growth' and component in custom_growth_rates:
                                avg_growth = custom_growth_rates[component]
                            else:
                                final_value = predictions[-1] if len(predictions) > 0 else data[component].iloc[-1]
                                avg_growth = ((final_value / data[component].iloc[-1]) ** (1/forecast_years) - 1) * 100
                            row['Avg Annual Growth (%)'] = f"{avg_growth:.2f}%"
                        else:
                            row['Avg Annual Growth (%)'] = None
                        
                        forecast_summary.append(row)
                
                if forecast_summary:
                    import pandas as pd
                    forecast_df = pd.DataFrame(forecast_summary)
                    st.dataframe(forecast_df, use_container_width=True)
                    
                    # Additional insights
                    st.markdown("**📊 Key Insights:**")
                    if len(forecast_summary) > 0:
                        # Find component with highest projected growth
                        growth_rates = []
                        for row in forecast_summary:
                            if row['Avg Annual Growth (%)'] and row['Avg Annual Growth (%)'] != 'N/A':
                                try:
                                    growth_val = float(row['Avg Annual Growth (%)'].replace('%', ''))
                                    growth_rates.append((row['Component'], growth_val))
                                except:
                                    continue
                        
                        if growth_rates:
                            highest_growth = max(growth_rates, key=lambda x: x[1])
                            lowest_growth = min(growth_rates, key=lambda x: x[1])
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.info(f"🚀 **Fastest Growing**: {highest_growth[0]} ({highest_growth[1]:.1f}% annually)")
                            with col2:
                                st.warning(f"📉 **Slowest Growing**: {lowest_growth[0]} ({lowest_growth[1]:.1f}% annually)")
            else:
                st.warning("No forecast data available for the selected components.")
        
        if tab7 is not None:
            with tab7:
                st.subheader("Forecast Accuracy Assessment")
                st.markdown("**🎯 Model performance evaluation using historical data backtesting**")
                
                if accuracy_results:
                    # Display accuracy metrics
                    fig_accuracy = visualizer.create_accuracy_comparison(accuracy_results)
                    st.plotly_chart(fig_accuracy, use_container_width=True)
                    
                    # Method comparison section
                    st.subheader("📈 Forecast Method Comparison")
                    
                    # Run accuracy test for all methods
                    with st.spinner("Comparing all forecast methods..."):
                        all_methods = ['linear_regression', 'polynomial', 'moving_average', 'random_forest', 'svm', 'var_model', 'gmm', 'cge_model']
                        method_comparison = {}
                        
                        for method in all_methods:
                            try:
                                method_accuracy = analyzer.calculate_forecast_accuracy(data, method=method)
                                if method_accuracy:
                                    # Calculate average MAPE across all components
                                    mape_values = [metrics['mape'] for metrics in method_accuracy.values() if 'mape' in metrics]
                                    rmse_values = [metrics['rmse'] for metrics in method_accuracy.values() if 'rmse' in metrics]
                                    mae_values = [metrics['mae'] for metrics in method_accuracy.values() if 'mae' in metrics]
                                    
                                    if mape_values:
                                        method_comparison[method] = {
                                            'Avg_MAPE': sum(mape_values) / len(mape_values),
                                            'Avg_RMSE': sum(rmse_values) / len(rmse_values) if rmse_values else 0,
                                            'Avg_MAE': sum(mae_values) / len(mae_values) if mae_values else 0
                                        }
                            except Exception:
                                continue
                    
                    if method_comparison:
                        comparison_data = []
                        for method, metrics in method_comparison.items():
                            method_name = {
                                'linear_regression': 'Linear Regression',
                                'polynomial': 'Polynomial Regression',
                                'moving_average': 'Moving Average',
                                'random_forest': 'Random Forest',
                                'svm': 'Support Vector Machine',
                                'var_model': 'Vector Autoregression',
                                'gmm': 'Generalized Method of Moments',
                                'cge_model': 'Computable General Equilibrium'
                            }.get(method, method.title())
                            
                            comparison_data.append({
                                'Method': method_name,
                                'Avg MAPE (%)': f"{metrics['Avg_MAPE']:.2f}",
                                'Avg RMSE': f"{metrics['Avg_RMSE']:.2f}",
                                'Avg MAE': f"{metrics['Avg_MAE']:.2f}"
                            })
                        
                        if comparison_data:
                            import pandas as pd
                            comparison_df = pd.DataFrame(comparison_data)
                            
                            # Color code based on performance
                            def color_mape(val):
                                try:
                                    mape_val = float(val)
                                    if mape_val < 10:
                                        return 'color: green; font-weight: bold'  # Excellent
                                    elif mape_val < 20:
                                        return 'color: blue; font-weight: bold'   # Good
                                    elif mape_val < 50:
                                        return 'color: orange; font-weight: bold' # Reasonable
                                    else:
                                        return 'color: red; font-weight: bold'    # Poor
                                except:
                                    return ''
                            
                            styled_comparison = comparison_df.style.applymap(color_mape, subset=['Avg MAPE (%)'])
                            st.dataframe(styled_comparison, use_container_width=True)
                            
                            # Find best method
                            best_method = min(method_comparison.items(), key=lambda x: x[1]['Avg_MAPE'])
                            st.success(f"🏆 **Best Performing Method**: {best_method[0].replace('_', ' ').title()} (MAPE: {best_method[1]['Avg_MAPE']:.2f}%)")
                    
                    # Accuracy summary table
                    st.subheader("Detailed Accuracy Metrics")
                    accuracy_summary = []
                    for component, metrics in accuracy_results.items():
                        component_name = {
                            'consumption': 'Consumption',
                            'investment': 'Investment', 
                            'government': 'Government',
                            'exports': 'Exports',
                            'imports': 'Imports',
                            'gdp': 'GDP'
                        }.get(component, component.title())
                        
                        accuracy_summary.append({
                            'Component': component_name,
                            'MAPE (%)': f"{metrics['mape']:.2f}",
                            'RMSE': f"{metrics['rmse']:.2f}",
                            'MAE': f"{metrics['mae']:.2f}",
                            'Method': metrics['method'].replace('_', ' ').title()
                        })
                    
                    if accuracy_summary:
                        import pandas as pd
                        accuracy_df = pd.DataFrame(accuracy_summary)
                        
                        # Color code MAPE values based on interpretation guide
                        def color_accuracy_metrics(val):
                            try:
                                if 'MAPE' in str(val):
                                    return ''
                                mape_val = float(val)
                                if mape_val < 10:
                                    return 'color: green; font-weight: bold'  # Excellent
                                elif mape_val < 20:
                                    return 'color: blue; font-weight: bold'   # Good
                                elif mape_val < 50:
                                    return 'color: orange; font-weight: bold' # Reasonable
                                else:
                                    return 'color: red; font-weight: bold'    # Poor
                            except:
                                return ''
                        
                        styled_accuracy = accuracy_df.style.applymap(color_accuracy_metrics, subset=['MAPE (%)'])
                        st.dataframe(styled_accuracy, use_container_width=True)
                        
                        # Interpretation guide
                        st.subheader("📊 Accuracy Interpretation Guide")
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.markdown("""
                            **MAPE (Mean Absolute Percentage Error)**
                            - < 10%: Excellent forecast
                            - 10-20%: Good forecast
                            - 20-50%: Reasonable forecast
                            - > 50%: Poor forecast
                            """)
                        
                        with col2:
                            st.markdown("""
                            **RMSE (Root Mean Square Error)**
                            - Lower values indicate better accuracy
                            - Penalizes larger errors more heavily
                            - Compare relative to data scale
                            """)
                        
                        with col3:
                            st.markdown("""
                            **MAE (Mean Absolute Error)**
                            - Average of absolute forecast errors
                            - More intuitive interpretation
                            - Less sensitive to outliers than RMSE
                            """)
                else:
                    st.warning("No accuracy assessment data available. This may occur if there's insufficient historical data for backtesting.")
        
        # VAR-specific analysis (Impulse Response and Variance Decomposition)
        if st.session_state.get('enable_forecasting', False) and st.session_state.get('forecast_method') == 'var_model':
            with st.expander("🔍 VAR Model Analysis - Impulse Response & Variance Decomposition", expanded=False):
                st.markdown("**Advanced Vector Autoregression (VAR) analysis showing economic interdependencies**")
                
                # Get VAR analysis results
                var_results = analyzer.get_var_analysis_results(data)
                
                if var_results:
                    st.success(f"✅ VAR model fitted with {var_results['lag_order']} lags")
                    
                    # Impulse Response Functions
                    st.subheader("📈 Impulse Response Functions")
                    st.markdown("**Shows how each economic component responds to shocks in other components**")
                    
                    # Select variables for impulse response
                    available_vars = list(data.columns)
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        shock_var = st.selectbox(
                            "Shock Variable (source)",
                            available_vars,
                            help="Variable that receives the initial shock"
                        )
                    
                    with col2:
                        response_var = st.selectbox(
                            "Response Variable (target)",
                            available_vars,
                            help="Variable whose response to the shock is measured"
                        )
                    
                    if shock_var and response_var:
                        try:
                            # Get impulse response data
                            irf = var_results['impulse_response']
                            shock_idx = available_vars.index(shock_var)
                            response_idx = available_vars.index(response_var)
                            
                            # Extract the specific impulse response
                            irf_data = irf.irfs[response_idx, shock_idx, :]
                            
                            # Create and display the plot
                            fig_irf = visualizer.create_impulse_response_plot(irf_data, shock_var, response_var)
                            st.plotly_chart(fig_irf, use_container_width=True)
                            
                            # Interpretation
                            max_response = max(abs(irf_data))
                            max_period = list(irf_data).index(max(irf_data, key=abs))
                            
                            st.info(f"📊 **Interpretation**: Maximum response of {max_response:.4f} occurs at period {max_period}")
                            
                        except Exception as e:
                            st.error("Unable to generate impulse response plot. This may occur with insufficient data or model fitting issues.")
                    
                    # Variance Decomposition
                    st.subheader("📊 Variance Decomposition")
                    st.markdown("**Shows the proportion of forecast error variance attributable to each variable**")
                    
                    # Select variable for variance decomposition
                    decomp_var = st.selectbox(
                        "Select variable for variance decomposition",
                        available_vars,
                        help="Variable whose forecast error variance will be decomposed"
                    )
                    
                    if decomp_var:
                        try:
                            # Get variance decomposition data
                            fevd = var_results['variance_decomposition']
                            var_idx = available_vars.index(decomp_var)
                            
                            # Extract variance decomposition for selected variable
                            decomp_data = fevd.decomp[:, var_idx, :]
                            
                            # Create and display the plot
                            fig_decomp = visualizer.create_variance_decomposition_plot(decomp_data, available_vars)
                            st.plotly_chart(fig_decomp, use_container_width=True)
                            
                            # Summary statistics
                            st.markdown("**📈 10-Period Ahead Decomposition Summary:**")
                            final_decomp = decomp_data[-1, :] * 100  # Convert to percentages
                            
                            decomp_summary = []
                            for i, var_name in enumerate(available_vars):
                                decomp_summary.append({
                                    'Variable': var_name.title(),
                                    'Contribution (%)': f"{final_decomp[i]:.1f}%"
                                })
                            
                            import pandas as pd
                            decomp_df = pd.DataFrame(decomp_summary)
                            st.dataframe(decomp_df, use_container_width=True)
                            
                        except Exception as e:
                            st.error("Unable to generate variance decomposition plot. This may occur with insufficient data or model fitting issues.")
                
                else:
                    st.warning("⚠️ VAR analysis requires sufficient historical data (minimum 10 observations) and may not be available for all datasets.")
    
    # Export functionality
    st.header("📥 Export Comprehensive Report")
    
    if results:
        st.markdown("**Generate a complete analysis report including all tabs and visualizations**")
        
        # Language selection
        col1, col2, col3 = st.columns(3)
        
        with col1:
            report_language = st.selectbox(
                "📋 Report Language",
                options=['english', 'indonesian'],
                format_func=lambda x: {
                    'english': '🇺🇸 English',
                    'indonesian': '🇮🇩 Bahasa Indonesia'
                }[x],
                help="Choose the language for the comprehensive report"
            )
        
        with col2:
            report_format = st.selectbox(
                "📄 Export Format",
                options=['pdf', 'docx'],
                format_func=lambda x: {
                    'pdf': '📄 PDF Document',
                    'docx': '📝 Word Document (DOCX)'
                }[x],
                help="Choose the format for the report export"
            )
        
        with col3:
            include_forecasts = st.checkbox(
                "Include Forecast Data",
                value=True,
                help="Include forecasting results and accuracy metrics in the report"
            )
        
        # Region name input
        region_name = st.text_input(
            "🌍 Region/Country Name",
            value="",
            placeholder="e.g., Indonesia, Malaysia, Thailand, etc.",
            help="Enter the name of the region or country being analyzed for personalized report"
        )
        
        # Export button
        if st.button("🚀 Generate Comprehensive Report", type="primary"):
            with st.spinner(f"Generating {report_format.upper()} report in {report_language.title()}..."):
                try:
                    # Ensure all required data is available for VAR method
                    if results.get('analysis_method') == 'var' and forecast_results is None:
                        st.warning("⚠️ VAR method requires forecast data to be included. Please enable 'Include Forecast Data' option.")
                    else:
                        report_buffer = create_comprehensive_report(
                            results=results,
                            data=data,
                            forecast_results=forecast_results if include_forecasts else None,
                            accuracy_results=accuracy_results if include_forecasts else None,
                            language=report_language,
                            format=report_format,
                            visualizer=visualizer,
                            region_name=region_name if region_name.strip() else "Region"
                        )
                        
                        if report_buffer:
                            # Generate filename with timestamp
                            from datetime import datetime
                            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                            filename = f"economic_analysis_report_{timestamp}.{report_format}"
                            
                            # Determine MIME type
                            mime_type = "application/pdf" if report_format == 'pdf' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            
                            st.download_button(
                                label=f"📥 Download {report_format.upper()} Report",
                                data=report_buffer,
                                file_name=filename,
                                mime=mime_type,
                                help=f"Download the comprehensive analysis report as {report_format.upper()}"
                            )
                            
                            st.success(f"✅ Report generated successfully! Click the download button above to save the {report_format.upper()} file.")
                            
                            # Report summary
                            st.info(f"""
                            📊 **Report Contents:**
                            - Economic Growth Analysis Results
                            - Regression Analysis with Statistical Significance
                            - Primary Growth Driver Identification ({results.get('analysis_method', 'correlation_contribution')})
                            - Correlation and Contribution Analysis
                            - Data Visualizations and Charts
                            {'- Forecast Results and Accuracy Metrics' if include_forecasts else ''}
                            - Methodology and Interpretation Guide
                            
                            🌐 **Language**: {report_language.title()}
                            📄 **Format**: {report_format.upper()}
                            """)
                        else:
                            st.error("❌ Failed to generate report. Please try again or contact support.")
                        
                except Exception as e:
                    st.error(f"❌ Error generating report: {str(e)}")
                    st.info("💡 This may occur due to insufficient data or formatting issues. Please ensure all analysis tabs have been properly loaded.")
    
    else:
        st.warning("⚠️ No analysis results available. Please upload data and run the analysis first to generate a report.")
    
    # Quick export options
    st.subheader("📋 Quick Export Options")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📊 Download CSV Data"):
            # Create downloadable report
            report_data = create_analysis_report(results, data)
            st.download_button(
                label="📁 Download CSV Report",
                data=report_data.to_csv(index=False),
                file_name="economic_analysis_report.csv",
                mime="text/csv"
            )
    
    with col2:
        if st.button("📈 Download Summary Statistics"):
            if 'summary_stats' in results:
                st.download_button(
                    label="📁 Download Summary CSV",
                    data=results['summary_stats'].to_csv(),
                    file_name="summary_statistics.csv",
                    mime="text/csv"
                )

def display_methodology():
    """Display the methodology and instructions"""
    
    st.header("🎯 About This Tool")
    
    st.markdown("""
    This economic analysis tool helps identify which types of expenditure are driving economic growth in a region 
    by analyzing GDP expenditure data across four main categories:
    
    - **Consumption-based expenditure**: Household and private consumption
    - **Investment-based expenditure**: Gross fixed capital formation and business investment
    - **Government spending**: Public sector expenditure
    - **Net Exports**: Export minus import balance
    """)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📋 Data Requirements")
        st.markdown("""
        Your CSV file should contain:
        - **Time series data** (monthly, quarterly, or annual)
        - **Date column** or time index
        - **Expenditure columns** for different GDP components
        - **Numerical values** in consistent units
        
        **Supported formats:**
        - Date formats: YYYY-MM-DD, MM/DD/YYYY, etc.
        - Numeric formats: With or without currency symbols
        - Missing values: Will be handled automatically
        """)
    
    with col2:
        st.subheader("🔬 Analysis Methods")
        st.markdown("""
        The tool employs several analytical techniques:
        - **Growth rate calculations** (year-over-year, period-over-period)
        - **Correlation analysis** between expenditure types and GDP
        - **Contribution analysis** to measure impact on total growth
        - **Regression analysis** for statistical significance
        - **Trend analysis** to identify patterns over time
        """)
    
    st.subheader("🚀 Getting Started")
    st.markdown("""
    1. **Upload your CSV file** using the file uploader in the sidebar
    2. **Map your columns** to the appropriate expenditure categories
    3. **Configure analysis settings** (time period, etc.)
    4. **Run the analysis** and explore the results
    5. **Export results** for further use or reporting
    """)

def display_detailed_methodology():
    """Display detailed methodology information"""
    
    st.markdown("""
    ## 📊 Analytical Methodology
    
    ### Growth Rate Calculations
    - **Year-over-Year Growth**: ((Current Period - Previous Period) / Previous Period) × 100
    - **Compound Annual Growth Rate (CAGR)**: ((Ending Value / Beginning Value)^(1/n) - 1) × 100
    - **Moving Average Growth**: Smoothed growth rates using rolling windows
    
    ### Contribution Analysis
    - **Absolute Contribution**: Change in expenditure component × weight in total GDP
    - **Relative Contribution**: (Component Growth × Component Share) / Total GDP Growth
    - **Marginal Contribution**: Incremental impact of each expenditure type
    
    ### Statistical Analysis
    - **Pearson Correlation**: Measures linear relationship strength
    - **Linear Regression**: Y = α + βX + ε, where Y is GDP growth and X is expenditure growth
    - **R-squared**: Proportion of variance explained by the model
    - **P-values**: Statistical significance testing (α = 0.05)
    
    ### Growth Driver Identification
    The primary growth driver is determined by:
    1. **Highest correlation** with GDP growth
    2. **Largest contribution** to total growth
    3. **Statistical significance** of the relationship
    4. **Consistency** over the analysis period
    
    ### Data Processing
    - **Missing value handling**: Forward fill, interpolation, or exclusion
    - **Outlier detection**: Statistical methods to identify anomalous values
    - **Seasonality adjustment**: For quarterly/monthly data
    - **Unit consistency**: Automatic scaling and normalization
    """)

def create_analysis_report(results, data):
    """Create a comprehensive analysis report for download"""
    
    report_rows = []
    
    # Add summary information
    report_rows.append({
        'Metric': 'Primary Growth Driver',
        'Value': results.get('primary_driver', 'N/A'),
        'Description': 'Expenditure category with highest contribution to GDP growth'
    })
    
    report_rows.append({
        'Metric': 'GDP Growth Rate (%)',
        'Value': f"{results.get('gdp_growth_rate', 0):.2f}",
        'Description': 'Average annual GDP growth rate'
    })
    
    report_rows.append({
        'Metric': 'Strongest Correlation',
        'Value': f"{results.get('strongest_correlation', 0):.3f}",
        'Description': 'Highest correlation between expenditure and GDP growth'
    })
    
    # Add correlation results
    if 'correlations' in results:
        for category, correlation in results['correlations'].items():
            if isinstance(correlation, (int, float)):
                report_rows.append({
                    'Metric': f'{category.title()} Correlation',
                    'Value': f"{correlation:.3f}",
                    'Description': f'Correlation between {category} and GDP growth'
                })
    
    # Add growth contributions
    if 'growth_contributions' in results:
        for category, contribution in results['growth_contributions'].items():
            if isinstance(contribution, (int, float)):
                report_rows.append({
                    'Metric': f'{category.title()} Contribution (%)',
                    'Value': f"{contribution:.2f}",
                    'Description': f'Percentage contribution of {category} to GDP growth'
                })
    
    return pd.DataFrame(report_rows)

def create_comprehensive_report(results, data, forecast_results=None, accuracy_results=None, language='english', format='pdf', visualizer=None, region_name="Region"):
    """Create comprehensive multilingual report in PDF or DOCX format"""
    
    try:
        # Language translations
        translations = {
            'english': {
                'title': f'Economic Growth Analysis Report - {region_name}',
                'executive_summary': 'Executive Summary',
                'data_overview': 'Data Overview',
                'regression_analysis': 'Regression Analysis',
                'primary_driver': 'Primary Growth Driver',
                'forecast_results': 'Forecast Results',
                'accuracy_metrics': 'Forecast Accuracy Metrics',
                'conclusions': 'Conclusions and Recommendations',
                'analysis_period': 'Analysis Period',
                'primary_driver_identified': 'Primary Driver Identified',
                'analysis_method': 'Analysis Method Used',
                'statistically_significant': 'Statistically Significant',
                'not_significant': 'Not Statistically Significant',
                'excellent_forecast': 'Excellent Forecast Accuracy',
                'good_forecast': 'Good Forecast Accuracy',
                'reasonable_forecast': 'Reasonable Forecast Accuracy',
                'poor_forecast': 'Poor Forecast Accuracy',
                'region_analyzed': 'Region/Country Analyzed'
            },
            'indonesian': {
                'title': f'Laporan Analisis Pertumbuhan Ekonomi - {region_name}',
                'executive_summary': 'Ringkasan Eksekutif',
                'data_overview': 'Tinjauan Data',
                'regression_analysis': 'Analisis Regresi',
                'primary_driver': 'Pendorong Pertumbuhan Utama',
                'forecast_results': 'Hasil Peramalan',
                'accuracy_metrics': 'Metrik Akurasi Peramalan',
                'conclusions': 'Kesimpulan dan Rekomendasi',
                'analysis_period': 'Periode Analisis',
                'primary_driver_identified': 'Pendorong Utama yang Teridentifikasi',
                'analysis_method': 'Metode Analisis yang Digunakan',
                'statistically_significant': 'Signifikan Secara Statistik',
                'not_significant': 'Tidak Signifikan Secara Statistik',
                'excellent_forecast': 'Akurasi Peramalan Sangat Baik',
                'good_forecast': 'Akurasi Peramalan Baik',
                'reasonable_forecast': 'Akurasi Peramalan Cukup Baik',
                'poor_forecast': 'Akurasi Peramalan Buruk',
                'region_analyzed': 'Wilayah/Negara yang Dianalisis'
            }
        }
        
        t = translations.get(language, translations['english'])
        
        if format == 'pdf':
            return create_pdf_report(results, data, forecast_results, accuracy_results, t, visualizer, region_name)
        else:
            return create_docx_report(results, data, forecast_results, accuracy_results, t, visualizer, region_name)
            
    except Exception as e:
        print(f"Error creating report: {e}")
        return None

def create_pdf_report(results, data, forecast_results, accuracy_results, translations, visualizer, region_name):
    """Create PDF report using ReportLab"""
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from io import BytesIO
        import datetime
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30, alignment=1)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceAfter=12, textColor=colors.darkblue)
        
        # Story elements
        story = []
        
        # Title
        title = Paragraph(translations['title'], title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Report metadata
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"Generated on: {current_date}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph(translations['executive_summary'], heading_style))
        
        primary_driver = results.get('primary_driver', 'Not determined')
        analysis_method = results.get('analysis_method', 'correlation_contribution')
        method_name = {
            'correlation_contribution': 'Correlation & Contribution Analysis',
            'log_log_regression': 'Log-Log Regression Analysis',
            'average_elasticity': 'Average Elasticity Method'
        }.get(analysis_method, analysis_method)
        
        summary_text = f"""
        This report presents a comprehensive analysis of economic growth drivers for <b>{region_name}</b> based on expenditure components. 
        The primary growth driver identified is: <b>{primary_driver}</b> using {method_name}.
        
        Analysis covers correlation analysis, contribution analysis, and regression modeling to determine 
        which expenditure components most significantly influence economic growth in {region_name}.
        """
        
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Data Overview
        story.append(Paragraph(translations['data_overview'], heading_style))
        
        data_info = f"""
        <b>{translations['region_analyzed']}:</b> {region_name}<br/>
        <b>{translations['analysis_period']}:</b> {data.index.min().strftime('%Y-%m-%d')} to {data.index.max().strftime('%Y-%m-%d')}<br/>
        <b>Total Observations:</b> {len(data)}<br/>
        <b>Components Analyzed:</b> {', '.join([col.title() for col in data.columns])}<br/>
        <b>{translations['analysis_method']}:</b> {method_name}
        """
        
        story.append(Paragraph(data_info, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Analysis Results based on method
        analysis_method = results.get('analysis_method', 'correlation_contribution')
        
        if analysis_method == 'log_log_regression' and 'log_log_elasticities' in results:
            story.append(Paragraph("Log-Log Regression Elasticity Analysis", heading_style))
            
            # Create elasticity table
            elasticity_data = [['Component', 'Elasticity', 'R-squared', 'P-value', 'Significance']]
            
            for component, elastic_result in results['log_log_elasticities'].items():
                component_name = component.title()
                elasticity = f"{elastic_result.get('elasticity', 0):.4f}"
                r_squared = f"{elastic_result.get('r_squared', 0):.4f}"
                p_value = f"{elastic_result.get('p_value', 1):.4f}"
                significance = translations['statistically_significant'] if elastic_result.get('p_value', 1) < 0.05 else translations['not_significant']
                
                elasticity_data.append([component_name, elasticity, r_squared, p_value, significance])
            
            elasticity_table = Table(elasticity_data)
            elasticity_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(elasticity_table)
            story.append(Spacer(1, 20))
            
        elif analysis_method == 'average_elasticity' and 'average_elasticities' in results:
            story.append(Paragraph("Average Elasticity Analysis", heading_style))
            
            # Create average elasticity table
            avg_elasticity_data = [['Component', 'Avg Elasticity', 'Abs Avg Elasticity', 'Std Dev', 'Correlation', 'Observations']]
            
            for component, elastic_result in results['average_elasticities'].items():
                component_name = component.title()
                avg_elasticity = f"{elastic_result.get('average_elasticity', 0):.4f}"
                abs_avg_elasticity = f"{elastic_result.get('absolute_average_elasticity', 0):.4f}"
                std_dev = f"{elastic_result.get('standard_deviation', 0):.4f}"
                correlation = f"{elastic_result.get('correlation', 0):.4f}"
                observations = str(elastic_result.get('observations', 0))
                
                avg_elasticity_data.append([component_name, avg_elasticity, abs_avg_elasticity, std_dev, correlation, observations])
            
            avg_elasticity_table = Table(avg_elasticity_data)
            avg_elasticity_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(avg_elasticity_table)
            story.append(Spacer(1, 20))
            
        elif 'regression_results' in results:  # Default correlation_contribution method
            story.append(Paragraph(translations['regression_analysis'], heading_style))
            
            # Create regression table
            regression_data = [['Component', 'R-squared', 'Coefficient', 'P-value', 'Significance']]
            
            for component, reg_result in results['regression_results'].items():
                component_name = component.title()
                r_squared = f"{reg_result.get('r_squared', 0):.4f}"
                coefficient = f"{reg_result.get('coefficient', 0):.4f}"
                p_value = f"{reg_result.get('p_value', 1):.4f}"
                significance = translations['statistically_significant'] if reg_result.get('p_value', 1) < 0.05 else translations['not_significant']
                
                regression_data.append([component_name, r_squared, coefficient, p_value, significance])
            
            regression_table = Table(regression_data)
            regression_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(regression_table)
            story.append(Spacer(1, 20))
        
        # Correlation and Contribution Analysis
        if 'correlations' in results:
            story.append(Paragraph("Correlation Analysis", heading_style))
            
            correlation_data = [['Component', 'Correlation with GDP']]
            for component, correlation in results['correlations'].items():
                if isinstance(correlation, (int, float)):
                    correlation_data.append([component.title(), f"{correlation:.4f}"])
            
            correlation_table = Table(correlation_data)
            correlation_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(correlation_table)
            story.append(Spacer(1, 20))
        
        if 'growth_contributions' in results:
            story.append(Paragraph("Growth Contribution Analysis", heading_style))
            
            contribution_data = [['Component', 'Contribution to GDP Growth (%)']]
            for component, contribution in results['growth_contributions'].items():
                if isinstance(contribution, (int, float)):
                    contribution_data.append([component.title(), f"{contribution:.2f}%"])
            
            contribution_table = Table(contribution_data)
            contribution_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(contribution_table)
            story.append(Spacer(1, 20))
        
        # Forecast Results (if available)
        if forecast_results:
            story.append(Paragraph(translations['forecast_results'], heading_style))
            forecast_text = f"Forecasting analysis was conducted for {region_name} using advanced econometric methods."
            story.append(Paragraph(forecast_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Accuracy metrics (if available)
            if accuracy_results:
                story.append(Paragraph(translations['accuracy_metrics'], heading_style))
                
                accuracy_data = [['Component', 'MAPE (%)', 'RMSE', 'MAE', 'Accuracy Level']]
                
                for component, metrics in accuracy_results.items():
                    component_name = component.title()
                    mape = f"{metrics['mape']:.2f}"
                    rmse = f"{metrics['rmse']:.2f}"
                    mae = f"{metrics['mae']:.2f}"
                    
                    # Determine accuracy level
                    mape_val = metrics['mape']
                    if mape_val < 10:
                        accuracy_level = translations['excellent_forecast']
                    elif mape_val < 20:
                        accuracy_level = translations['good_forecast']
                    elif mape_val < 50:
                        accuracy_level = translations['reasonable_forecast']
                    else:
                        accuracy_level = translations['poor_forecast']
                    
                    accuracy_data.append([component_name, mape, rmse, mae, accuracy_level])
                
                accuracy_table = Table(accuracy_data)
                accuracy_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(accuracy_table)
                story.append(Spacer(1, 20))
        
        # Conclusions
        story.append(Paragraph(translations['conclusions'], heading_style))
        
        conclusions_text = f"""
        Based on the comprehensive analysis conducted for {region_name}:
        
        1. <b>{translations['primary_driver_identified']}:</b> {primary_driver}
        2. The analysis utilized {method_name} to determine growth drivers
        3. Statistical significance was evaluated using p-values with α = 0.05 threshold
        4. Results provide actionable insights for economic policy and planning in {region_name}
        """
        
        story.append(Paragraph(conclusions_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return None

def create_docx_report(results, data, forecast_results, accuracy_results, translations, visualizer, region_name):
    """Create DOCX report using python-docx"""
    
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from io import BytesIO
        import datetime
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(translations['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report metadata
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph(f"Generated on: {current_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading(translations['executive_summary'], level=1)
        
        primary_driver = results.get('primary_driver', 'Not determined')
        analysis_method = results.get('analysis_method', 'correlation_contribution')
        method_name = {
            'correlation_contribution': 'Correlation & Contribution Analysis',
            'log_log_regression': 'Log-Log Regression Analysis', 
            'average_elasticity': 'Average Elasticity Method'
        }.get(analysis_method, analysis_method)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"This report presents a comprehensive analysis of economic growth drivers for ")
        summary_para.add_run(region_name).bold = True
        summary_para.add_run(f" based on expenditure components. The primary growth driver identified is: ")
        summary_para.add_run(primary_driver).bold = True
        summary_para.add_run(f" using {method_name}.")
        
        # Data Overview
        doc.add_heading(translations['data_overview'], level=1)
        
        data_table = doc.add_table(rows=5, cols=2)
        data_table.style = 'Table Grid'
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        data_cells = [
            [translations['region_analyzed'], region_name],
            [translations['analysis_period'], f"{data.index.min().strftime('%Y-%m-%d')} to {data.index.max().strftime('%Y-%m-%d')}"],
            ["Total Observations", str(len(data))],
            ["Components Analyzed", ', '.join([col.title() for col in data.columns])],
            [translations['analysis_method'], method_name]
        ]
        
        for i, (key, value) in enumerate(data_cells):
            data_table.cell(i, 0).text = key
            data_table.cell(i, 1).text = value
            # Make first column bold
            data_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Analysis Results based on method
        analysis_method = results.get('analysis_method', 'correlation_contribution')
        
        if analysis_method == 'log_log_regression' and 'log_log_elasticities' in results:
            doc.add_heading("Log-Log Regression Elasticity Analysis", level=1)
            
            # Create elasticity table
            elasticity_table = doc.add_table(rows=1, cols=5)
            elasticity_table.style = 'Table Grid'
            elasticity_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            headers = ['Component', 'Elasticity', 'R-squared', 'P-value', 'Significance']
            header_cells = elasticity_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for component, elastic_result in results['log_log_elasticities'].items():
                row_cells = elasticity_table.add_row().cells
                row_cells[0].text = component.title()
                row_cells[1].text = f"{elastic_result.get('elasticity', 0):.4f}"
                row_cells[2].text = f"{elastic_result.get('r_squared', 0):.4f}"
                row_cells[3].text = f"{elastic_result.get('p_value', 1):.4f}"
                row_cells[4].text = translations['statistically_significant'] if elastic_result.get('p_value', 1) < 0.05 else translations['not_significant']
                
        elif analysis_method == 'average_elasticity' and 'average_elasticities' in results:
            doc.add_heading("Average Elasticity Analysis", level=1)
            
            # Create average elasticity table
            avg_elasticity_table = doc.add_table(rows=1, cols=6)
            avg_elasticity_table.style = 'Table Grid'
            avg_elasticity_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            headers = ['Component', 'Avg Elasticity', 'Abs Avg Elasticity', 'Std Dev', 'Correlation', 'Observations']
            header_cells = avg_elasticity_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for component, elastic_result in results['average_elasticities'].items():
                row_cells = avg_elasticity_table.add_row().cells
                row_cells[0].text = component.title()
                row_cells[1].text = f"{elastic_result.get('average_elasticity', 0):.4f}"
                row_cells[2].text = f"{elastic_result.get('absolute_average_elasticity', 0):.4f}"
                row_cells[3].text = f"{elastic_result.get('standard_deviation', 0):.4f}"
                row_cells[4].text = f"{elastic_result.get('correlation', 0):.4f}"
                row_cells[5].text = str(elastic_result.get('observations', 0))
                
        elif 'regression_results' in results:  # Default correlation_contribution method
            doc.add_heading(translations['regression_analysis'], level=1)
            
            # Create regression table
            reg_table = doc.add_table(rows=1, cols=5)
            reg_table.style = 'Table Grid'
            reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            headers = ['Component', 'R-squared', 'Coefficient', 'P-value', 'Significance']
            header_cells = reg_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for component, reg_result in results['regression_results'].items():
                row_cells = reg_table.add_row().cells
                row_cells[0].text = component.title()
                row_cells[1].text = f"{reg_result.get('r_squared', 0):.4f}"
                row_cells[2].text = f"{reg_result.get('coefficient', 0):.4f}"
                row_cells[3].text = f"{reg_result.get('p_value', 1):.4f}"
                row_cells[4].text = translations['statistically_significant'] if reg_result.get('p_value', 1) < 0.05 else translations['not_significant']
        
        # Correlation Analysis
        if 'correlations' in results:
            doc.add_heading("Correlation Analysis", level=1)
            
            corr_table = doc.add_table(rows=1, cols=2)
            corr_table.style = 'Table Grid'
            corr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            corr_headers = ['Component', 'Correlation with GDP']
            corr_header_cells = corr_table.rows[0].cells
            for i, header in enumerate(corr_headers):
                corr_header_cells[i].text = header
                corr_header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for component, correlation in results['correlations'].items():
                if isinstance(correlation, (int, float)):
                    row_cells = corr_table.add_row().cells
                    row_cells[0].text = component.title()
                    row_cells[1].text = f"{correlation:.4f}"
        
        # Growth Contribution Analysis
        if 'growth_contributions' in results:
            doc.add_heading("Growth Contribution Analysis", level=1)
            
            contrib_table = doc.add_table(rows=1, cols=2)
            contrib_table.style = 'Table Grid'
            contrib_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            contrib_headers = ['Component', 'Contribution to GDP Growth (%)']
            contrib_header_cells = contrib_table.rows[0].cells
            for i, header in enumerate(contrib_headers):
                contrib_header_cells[i].text = header
                contrib_header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for component, contribution in results['growth_contributions'].items():
                if isinstance(contribution, (int, float)):
                    row_cells = contrib_table.add_row().cells
                    row_cells[0].text = component.title()
                    row_cells[1].text = f"{contribution:.2f}%"
        
        # Forecast Results (if available)
        if forecast_results:
            doc.add_heading(translations['forecast_results'], level=1)
            doc.add_paragraph(f"Forecasting analysis was conducted for {region_name} using advanced econometric methods.")
            
            # Accuracy metrics (if available)
            if accuracy_results:
                doc.add_heading(translations['accuracy_metrics'], level=2)
                
                acc_table = doc.add_table(rows=1, cols=5)
                acc_table.style = 'Table Grid'
                acc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Headers
                acc_headers = ['Component', 'MAPE (%)', 'RMSE', 'MAE', 'Accuracy Level']
                acc_header_cells = acc_table.rows[0].cells
                for i, header in enumerate(acc_headers):
                    acc_header_cells[i].text = header
                    acc_header_cells[i].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for component, metrics in accuracy_results.items():
                    row_cells = acc_table.add_row().cells
                    row_cells[0].text = component.title()
                    row_cells[1].text = f"{metrics['mape']:.2f}"
                    row_cells[2].text = f"{metrics['rmse']:.2f}"
                    row_cells[3].text = f"{metrics['mae']:.2f}"
                    
                    # Determine accuracy level
                    mape_val = metrics['mape']
                    if mape_val < 10:
                        accuracy_level = translations['excellent_forecast']
                    elif mape_val < 20:
                        accuracy_level = translations['good_forecast']
                    elif mape_val < 50:
                        accuracy_level = translations['reasonable_forecast']
                    else:
                        accuracy_level = translations['poor_forecast']
                    
                    row_cells[4].text = accuracy_level
        
        # Conclusions
        doc.add_heading(translations['conclusions'], level=1)
        
        conclusions = [
            f"{translations['primary_driver_identified']}: {primary_driver}",
            f"The analysis utilized {method_name} to determine growth drivers",
            "Statistical significance was evaluated using p-values with α = 0.05 threshold",
            f"Results provide actionable insights for economic policy and planning in {region_name}"
        ]
        
        for i, conclusion in enumerate(conclusions, 1):
            para = doc.add_paragraph(f"{i}. ")
            if i == 1:
                para.add_run(translations['primary_driver_identified']).bold = True
                para.add_run(f": {primary_driver}")
            else:
                para.add_run(conclusion)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return None

if __name__ == "__main__":
    main()
